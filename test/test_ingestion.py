"""
Formal test suite for ingestion.py - text extraction, content hashing, and
chunking. Every fixture used here is generated fresh in a tmp_path, nothing
depends on files existing on disk outside the test run.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.ingestion import (
    IngestionError,
    chunk_document,
    chunk_text,
    compute_file_hash,
    extract_pages,
)


# ---------------------------------------------------------------------------
# Fixtures - real files, generated fresh per test run
# ---------------------------------------------------------------------------

@pytest.fixture
def txt_file(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("This is a simple text document. " * 50, encoding="utf-8")
    return path


@pytest.fixture
def empty_txt_file(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n  \n", encoding="utf-8")
    return path


@pytest.fixture
def docx_file(tmp_path):
    path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Test Document", level=1)
    for i in range(10):
        doc.add_paragraph(f"This is paragraph {i} with some real content in it for testing purposes.")
    doc.save(str(path))
    return path


@pytest.fixture
def pdf_file(tmp_path):
    """A real 3-page PDF, so page-number tracking can be tested for real."""
    path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(path), pagesize=LETTER)
    for page_num in range(3):
        c.drawString(72, 750, f"Page {page_num + 1} heading")
        for line in range(20):
            c.drawString(72, 720 - line * 20, f"Page {page_num + 1} content line {line} about topic {page_num}.")
        c.showPage()
    c.save()
    return path


# ---------------------------------------------------------------------------
# extract_pages
# ---------------------------------------------------------------------------

def test_extract_pages_txt_returns_single_page_with_no_page_number(txt_file):
    pages = extract_pages(txt_file)
    assert len(pages) == 1
    page_number, text = pages[0]
    assert page_number is None
    assert "simple text document" in text


def test_extract_pages_empty_txt_raises(empty_txt_file):
    with pytest.raises(IngestionError):
        extract_pages(empty_txt_file)


def test_extract_pages_docx_returns_single_page_with_no_page_number(docx_file):
    pages = extract_pages(docx_file)
    assert len(pages) == 1
    page_number, text = pages[0]
    assert page_number is None
    assert "paragraph 0" in text
    assert "paragraph 9" in text


def test_extract_pages_pdf_returns_one_entry_per_page_with_correct_numbers(pdf_file):
    pages = extract_pages(pdf_file)
    assert len(pages) == 3
    page_numbers = [p[0] for p in pages]
    assert page_numbers == [1, 2, 3], "PDF pages should be 1-indexed and in order"
    assert "Page 1" in pages[0][1]
    assert "Page 2" in pages[1][1]
    assert "Page 3" in pages[2][1]


def test_extract_pages_unsupported_extension_raises(tmp_path):
    bad_file = tmp_path / "sample.xyz"
    bad_file.write_text("content")
    with pytest.raises(IngestionError, match="Unsupported file type"):
        extract_pages(bad_file)


# ---------------------------------------------------------------------------
# compute_file_hash
# ---------------------------------------------------------------------------

def test_compute_file_hash_is_deterministic(tmp_path):
    f1 = tmp_path / "a.txt"
    f2 = tmp_path / "b.txt"
    f1.write_text("identical content")
    f2.write_text("identical content")
    assert compute_file_hash(f1) == compute_file_hash(f2)


def test_compute_file_hash_differs_for_different_content(tmp_path):
    f1 = tmp_path / "a.txt"
    f2 = tmp_path / "b.txt"
    f1.write_text("content A")
    f2.write_text("content B")
    assert compute_file_hash(f1) != compute_file_hash(f2)


def test_compute_file_hash_handles_large_files_via_chunked_reads(tmp_path):
    # Larger than the 8192-byte read block used internally, to exercise the
    # actual chunked-reading loop rather than a single read() call.
    f = tmp_path / "large.txt"
    f.write_text("x" * 50_000)
    h = compute_file_hash(f)
    assert len(h) == 64  # SHA256 hex digest length


# ---------------------------------------------------------------------------
# chunk_text
# ---------------------------------------------------------------------------

def test_chunk_text_respects_target_size():
    text = " ".join(f"word{i}" for i in range(100))
    chunks = chunk_text(text, chunk_size_words=20, overlap_words=0)
    assert len(chunks) == 5
    for c in chunks:
        assert len(c.split()) == 20


def test_chunk_text_overlap_duplicates_boundary_words():
    text = " ".join(f"word{i}" for i in range(1, 21))  # word1..word20
    chunks = chunk_text(text, chunk_size_words=10, overlap_words=3)
    # chunk 0: word1-word10, chunk 1 should start with the last 3 words of chunk 0
    chunk0_words = chunks[0].split()
    chunk1_words = chunks[1].split()
    assert chunk0_words[-3:] == chunk1_words[:3]


def test_chunk_text_empty_string_returns_no_chunks():
    assert chunk_text("", chunk_size_words=10, overlap_words=2) == []


def test_chunk_text_overlap_must_be_smaller_than_size():
    with pytest.raises(ValueError):
        chunk_text("some words here", chunk_size_words=5, overlap_words=5)


def test_chunk_text_full_coverage_no_word_dropped():
    """Every word from the source should appear somewhere in the chunks -
    overlap duplicates words, but chunking should never silently drop one."""
    words = [f"word{i}" for i in range(1, 51)]
    text = " ".join(words)
    chunks = chunk_text(text, chunk_size_words=10, overlap_words=2)
    all_chunked_words = set(" ".join(chunks).split())
    assert all_chunked_words == set(words)


# ---------------------------------------------------------------------------
# chunk_document — the full pipeline
# ---------------------------------------------------------------------------

def test_chunk_document_produces_chunks_with_correct_metadata(txt_file):
    doc_id, chunks = chunk_document(txt_file, doc_name="my_doc.txt")

    assert len(doc_id) == 12  # uuid4().hex[:12]
    assert len(chunks) > 0

    for i, chunk in enumerate(chunks):
        assert chunk.doc_id == doc_id
        assert chunk.doc_name == "my_doc.txt"
        assert chunk.chunk_index == i
        assert chunk.chunk_id == f"{doc_id}::chunk::{i}"
        assert chunk.content_hash == compute_file_hash(txt_file)
        assert chunk.stored_filename == txt_file.name


def test_chunk_document_defaults_doc_name_to_filename(txt_file):
    _, chunks = chunk_document(txt_file)  # no doc_name passed
    assert chunks[0].doc_name == txt_file.name


def test_chunk_document_preserves_pdf_page_numbers(pdf_file):
    _, chunks = chunk_document(pdf_file)
    page_numbers = {c.page_number for c in chunks}
    assert page_numbers == {1, 2, 3}
    # every chunk should belong to exactly one page - never spans two
    for c in chunks:
        assert c.page_number is not None


def test_chunk_document_txt_has_no_page_numbers(txt_file):
    _, chunks = chunk_document(txt_file)
    assert all(c.page_number is None for c in chunks)


def test_chunk_document_raises_on_unsupported_type(tmp_path):
    bad_file = tmp_path / "sample.xyz"
    bad_file.write_text("content")
    with pytest.raises(IngestionError):
        chunk_document(bad_file)


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"])) #first test passed with 18 items